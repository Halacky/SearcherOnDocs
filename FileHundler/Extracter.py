import pandas as ps
import FileHundler.OptimizeText as OptimizeText
import fitz
import docx
import datetime


## Функция получения текста из эксельки
def readTextExcel (file):
    try:
        print("Обрабатывается файл: " + file)
        text = ''
        excel_file = ps.read_excel(file, sheet_name=None, header=None, usecols=[x for x in range(30)], nrows=10000)  # При чтении в параметрах обрезаем эксель и читаем все листы
        for df in excel_file:
            data = excel_file[df]
            for index, row in data.iterrows():
                for col in data.columns:
                    if not ps.isnull(row[col]):
                        text += f"{row[col]} "

        optimized_text = OptimizeText.optimazeText(text)
        #return optimized_text
        if len(optimized_text) > 1:
            return optimized_text
        else:
            return False
    except Exception as e:
        return False

## Функция получения текста из PDF
def readTextPdf(file):
    try:
        print("Обрабатывается файл: " + file)
        text = ''
        doc = fitz.open(file)
        for page in range(doc.pageCount):
            print('==============')
            print(str(doc[page].getText()))
            print('==============')
            text += str(doc[page].getText()) + " "

        optimized_text = OptimizeText.optimazeText(text)
        #return optimized_text
        if len(optimized_text) > 100:
            return optimized_text
        else:
            return False

    except Exception as e:
        print(e)
        return False


# Функция получения текста из ворда
def readTextWord(file):
    text = ''
    try:
        print("Обрабатывается файл: " + file)
        file_word = docx.Document(file)
        text = '/n'.join([p.text+" " for p in file_word.paragraphs])  # Получаем все праграфы
        text += get_textTable(file_word)
        optimized_text = OptimizeText.optimazeText(text)
        print("=============Optimized")
        #return optimized_text
        if len(optimized_text) > 1:
            print("=============Optimized (len)")
            return optimized_text
        else:
            return False
    except Exception as e:
        return False


## Метод для чтения текста из таблиц в ворде
def get_textTable(doc):
    text_table =[]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_table.append(cell.text)
                print("======Time: " + str(datetime.datetime.now()) + " Text=======" + str(cell.text) + "\n")

    return '/n'+' '.join(text_table)


## Метод для чтения текста из текстового файла
def readTextTXT(file):
    text = ""
    try:
        print("Обрабатывается файл: " + file)
        with open(file, "r", encoding="utf-8") as file:
            text = " ".join(file.readlines())
        optimized_text = OptimizeText.optimazeText(text)
        print("=============Optimized")
        #return optimized_text
        if len(optimized_text) > 1:
            print("=============Optimized (len)")
            return optimized_text
        else:
            return False
    except Exception as e:
        return False